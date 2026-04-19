"""DOCX report generator using python-docx."""

from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt, RGBColor

from reporter.common import compute_forecast


def build_docx_report(filepath: str, route_id: int, model_type: str, horizon: int) -> str:
    data = compute_forecast(route_id, model_type, horizon)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Отчёт о прогнозе пассажиропотока")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Сгенерировано: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x7f, 0x7f, 0x7f)

    doc.add_paragraph()

    # Параметры
    doc.add_heading("1. Параметры прогноза", level=2)
    params_table = doc.add_table(rows=0, cols=2)
    params_table.style = "Light Grid Accent 1"
    for k, v in [
        ("Маршрут",                data["route_code"]),
        ("Модель",                 model_type.upper()),
        ("Горизонт, мес.",         str(horizon)),
        ("Обучающих наблюдений",   str(len(data["history"]))),
    ]:
        row = params_table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_paragraph()

    # Метрики
    doc.add_heading("2. Метрики качества на тестовой выборке", level=2)
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = "Light Grid Accent 1"
    hdr = metrics_table.rows[0].cells
    hdr[0].text = "Метрика"
    hdr[1].text = "Значение"
    metrics = data["metrics"]
    for k, v in [
        ("MAPE, %",     f"{metrics['mape']:.2f}"),
        ("MAE, пасс.",  f"{metrics['mae']:.0f}"),
        ("RMSE, пасс.", f"{metrics['rmse']:.0f}"),
    ]:
        row = metrics_table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_paragraph()

    # Прогнозные значения
    doc.add_heading("3. Прогнозные значения с 95% ДИ", level=2)
    fc_table = doc.add_table(rows=1, cols=5)
    fc_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Месяц", "Прогноз", "Нижняя ДИ", "Верхняя ДИ", "Факт"]):
        fc_table.rows[0].cells[i].text = h
    for i, (p_v, l, u, a) in enumerate(zip(data["forecast"], data["lower"], data["upper"], data["actual"])):
        row = fc_table.add_row().cells
        row[0].text = f"+{i+1}"
        row[1].text = f"{p_v:,.0f}"
        row[2].text = f"{l:,.0f}"
        row[3].text = f"{u:,.0f}"
        row[4].text = f"{a:,.0f}"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    r = footer.add_run(
        "Отчёт сформирован автоматически информационной системой прогнозирования "
        "пассажиропотока на междугородних автобусных рейсах (МТИ, 2026).")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x7f, 0x7f, 0x7f)

    doc.save(filepath)
    return filepath
